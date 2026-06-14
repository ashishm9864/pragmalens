import html
import io
import os
import re
from statistics import mean

import pandas as pd
import plotly.express as px
import streamlit as st

from examples import ARTICLE_EXAMPLES, COMPARE_PAIRS, EVALUATION_CASES, EXAMPLES, KILLER_DEMOS
from nlp_pipeline import load_spacy_model, run_pipeline
from trigger_rules import TRIGGER_COLORS, TRIGGER_LABELS
from utils import (
    build_json_report,
    build_markdown_report,
    compute_media_literacy_score,
    display_confidence_chart,
    display_highlighted_text,
    display_legend,
    display_media_literacy_score,
    display_pipeline_visual,
    display_presupposition_card,
    display_rule_evidence,
    display_trigger_chart,
    negation_probe_text,
    presuppositions_to_rows,
)


st.set_page_config(page_title="PragmaLens", page_icon="PL", layout="wide")

st.markdown(
    """
    <style>
    :root {
        --navy: #0F172A;
        --muted: #475569;
        --line: #E2E8F0;
        --panel: #F8FAFC;
        --accent: #2563EB;
    }
    html, body, [class*="css"] {
        font-family: Inter, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
        color: var(--navy);
    }
    .block-container {
        padding-top: 2.2rem;
        padding-bottom: 2.5rem;
        max-width: 1220px;
    }
    h1, h2, h3 {
        letter-spacing: 0;
    }
    .pl-header {
        border-bottom: 1px solid var(--line);
        padding-bottom: 1rem;
        margin-bottom: 1.2rem;
    }
    .pl-caption {
        color: var(--muted);
        font-size: 1.05rem;
        margin-top: -.6rem;
    }
    .pl-callout {
        background: var(--panel);
        border: 1px solid var(--line);
        border-radius: 8px;
        padding: 15px 17px;
        color: #334155;
        line-height: 1.55;
        min-height: 100%;
    }
    .pl-tag {
        display: inline-block;
        border: 1px solid var(--line);
        border-radius: 6px;
        padding: 4px 8px;
        margin: 3px 4px 3px 0;
        font-size: 12px;
        color: #334155;
        background: #FFFFFF;
    }
    .persona-card {
        background: #F8FAFC;
        border: 1px solid #E2E8F0;
        border-radius: 12px;
        padding: 20px;
        height: 100%;
    }
    .persona-card h4 { margin-top: 0; color: #0F172A; }
    .persona-card .example-use {
        background: #EFF6FF;
        border-radius: 6px;
        padding: 8px 12px;
        font-size: 0.85rem;
        color: #2563EB;
        margin-top: 12px;
    }
    .hero-banner {
        background: linear-gradient(135deg, #EFF6FF 0%, #F5F3FF 100%);
        border-radius: 16px;
        padding: 28px 32px;
        margin-bottom: 24px;
    }
    .hero-banner h2 { font-size: 1.9rem; color: #0F172A; margin-bottom: 8px; }
    .hero-banner p { font-size: 1.05rem; color: #475569; margin: 0; }
    .pill-badge {
        display: inline-block;
        background: white;
        border: 1px solid #E2E8F0;
        border-radius: 20px;
        padding: 5px 14px;
        font-size: 0.82rem;
        color: #475569;
        margin: 12px 8px 0 0;
    }
    .score-card {
        border-radius: 12px;
        padding: 20px 22px;
        background: #F8FAFC;
        border-left: 5px solid;
        margin: 12px 0;
    }
    .score-number { font-size: 2.5rem; font-weight: 900; line-height: 1; }
    .score-bar-track {
        background: #E2E8F0;
        border-radius: 4px;
        height: 8px;
        margin: 10px 0;
    }
    .score-bar-fill {
        border-radius: 4px;
        height: 8px;
    }
    .compare-presup-chip {
        display: inline-block;
        background: #F0FDF4;
        border: 1px solid #86EFAC;
        border-radius: 6px;
        padding: 5px 10px;
        font-size: 0.82rem;
        color: #166534;
        margin: 4px 0;
    }
    .workflow-card {
        background: #FFFFFF;
        border: 1px solid #E2E8F0;
        border-radius: 10px;
        padding: 14px 16px;
        height: 100%;
    }
    .workflow-card strong {
        color: #0F172A;
    }
    .risk-sentence {
        border: 1px solid #E2E8F0;
        border-left: 5px solid #EF4444;
        border-radius: 10px;
        padding: 13px 15px;
        background: #FFFFFF;
        margin: 10px 0;
    }
    .stButton > button, .stDownloadButton > button {
        border-radius: 8px;
        font-weight: 700;
    }
    footer { visibility: hidden; }
    </style>
    """,
    unsafe_allow_html=True,
)


def has_api_key() -> bool:
    return any(
        os.environ.get(name)
        for name in ("GROQ_API_KEY", "OPENAI_API_KEY", "ANTHROPIC_API_KEY")
    )


def source_label(use_llm: bool) -> str:
    if not use_llm:
        return "Rule-based"
    if os.environ.get("GROQ_API_KEY"):
        return "Groq LLM"
    if os.environ.get("OPENAI_API_KEY"):
        return "OpenAI LLM"
    if os.environ.get("ANTHROPIC_API_KEY"):
        return "Anthropic LLM"
    return "Rule-based fallback"


def display_result_summary(results) -> None:
    if not results:
        return
    cols = st.columns(4)
    cols[0].metric("Detected", len(results))
    cols[1].metric("Avg confidence", f"{mean(item.confidence for item in results):.0%}")
    cols[2].metric("Avg subtlety", f"{mean(item.subtlety for item in results):.1f}/5")
    cols[3].metric("Trigger types", len({item.trigger_type for item in results}))


def display_downloads(sentence: str, results) -> None:
    col_md, col_json = st.columns(2)
    col_md.download_button(
        "Download Markdown report",
        build_markdown_report(sentence, results),
        file_name="pragmalens-analysis.md",
        mime="text/markdown",
        width="stretch",
    )
    col_json.download_button(
        "Download JSON report",
        build_json_report(sentence, results),
        file_name="pragmalens-analysis.json",
        mime="application/json",
        width="stretch",
    )


TRIGGER_SAMPLES = {
    "factive": "know, realize, discover...",
    "implicative": "manage, fail, forget...",
    "change_of_state": "stop, start, continue...",
    "iterative": "again, still, back...",
    "definite_np": "the report, the crisis...",
    "temporal": "before, after, when...",
    "cleft": "it was, what...",
}


def display_compact_presupposition(item) -> None:
    color = TRIGGER_COLORS.get(item.trigger_type, "#64748B")
    label = html.escape(TRIGGER_LABELS.get(item.trigger_type, item.trigger_type))
    presupposition = html.escape(item.presupposition_str)
    st.markdown(
        f"""
        <div style="border:1px solid #E2E8F0;border-radius:8px;padding:10px 12px;margin:8px 0;background:#FFFFFF;">
            <span style="display:inline-block;background:{color}18;color:{color};border:1px solid {color}55;
                         border-radius:999px;padding:2px 8px;font-size:0.75rem;font-weight:800;">
                {label}
            </span>
            <div style="color:#0F172A;margin-top:7px;line-height:1.45;">"{presupposition}"</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def run_evaluation() -> pd.DataFrame:
    rows = []
    for case in EVALUATION_CASES:
        results = run_pipeline(case["sentence"], use_llm=False)
        detected = {item.trigger_type for item in results}
        expected = case["expected"]
        passed = expected.issubset(detected)
        rows.append(
            {
                "sentence": case["sentence"],
                "expected": ", ".join(sorted(expected)),
                "detected": ", ".join(sorted(detected)) or "none",
                "pass": passed,
                "count": len(results),
            }
        )
    return pd.DataFrame(rows)


def extract_text_from_upload(uploaded_file) -> str:
    filename = uploaded_file.name.lower()
    raw = uploaded_file.getvalue()

    if filename.endswith(".docx"):
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX upload requires python-docx. Run: pip install python-docx") from exc

        document = Document(io.BytesIO(raw))
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("latin-1")


def split_text_into_sentences(text: str, limit: int) -> list[str]:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return []

    nlp = load_spacy_model()
    doc = nlp(cleaned[:200_000])
    sentences = [
        sentence.text.strip()
        for sentence in doc.sents
        if len(sentence.text.strip().split()) >= 4
    ]
    return sentences[:limit]


def analyze_sentence_collection(sentences: list[str], use_llm: bool, progress=None):
    summary_rows = []
    detail_rows = []
    sentence_results = []

    if not sentences:
        return pd.DataFrame(), pd.DataFrame(), sentence_results

    for index, sentence in enumerate(sentences, start=1):
        results = run_pipeline(sentence, use_llm=use_llm)
        score_dict = compute_media_literacy_score(results)
        media_score = score_dict["score"]
        trigger_types = sorted({TRIGGER_LABELS.get(item.trigger_type, item.trigger_type) for item in results})

        summary_rows.append(
            {
                "sentence_id": index,
                "sentence": sentence,
                "count": len(results),
                "media_literacy_score": media_score,
                "score_label": score_dict["label"],
                "avg_confidence": round(mean(item.confidence for item in results), 2) if results else 0,
                "avg_subtlety": round(mean(item.subtlety for item in results), 2) if results else 0,
                "trigger_types": ", ".join(trigger_types),
            }
        )
        for row in presuppositions_to_rows(sentence, results):
            row["sentence_id"] = index
            row["media_literacy_score"] = media_score
            detail_rows.append(row)
        sentence_results.append(
            {
                "sentence_id": index,
                "sentence": sentence,
                "results": results,
                "score": media_score,
                "score_label": score_dict["label"],
            }
        )
        if progress is not None:
            progress.progress(index / len(sentences))

    return pd.DataFrame(summary_rows), pd.DataFrame(detail_rows), sentence_results


def build_document_audit_report(title: str, summary_df: pd.DataFrame, detail_df: pd.DataFrame) -> str:
    if summary_df.empty:
        return "# PragmaLens Article Audit\n\nNo analyzable sentences were found."

    total_presuppositions = int(summary_df["count"].sum())
    avg_score = round(summary_df["media_literacy_score"].mean(), 1)
    high_load_count = int((summary_df["media_literacy_score"] > 6).sum())
    top_df = summary_df.sort_values(
        ["media_literacy_score", "count", "avg_subtlety"],
        ascending=False,
    ).head(8)

    lines = [
        "# PragmaLens Article Audit",
        "",
        f"**Document:** {title or 'Untitled draft'}",
        f"**Sentences scanned:** {len(summary_df)}",
        f"**Total presuppositions:** {total_presuppositions}",
        f"**Average manipulation load:** {avg_score}/10",
        f"**High-load sentences:** {high_load_count}",
        "",
        "## Highest-Risk Sentences",
        "",
    ]

    for _, row in top_df.iterrows():
        lines.extend(
            [
                f"### Sentence {int(row['sentence_id'])}: {row['media_literacy_score']}/10",
                row["sentence"],
                "",
                f"- Presupposition count: {int(row['count'])}",
                f"- Trigger types: {row['trigger_types'] or 'none'}",
                "",
            ]
        )

    if not detail_df.empty:
        lines.extend(["## Assumption Ledger", ""])
        for _, row in detail_df.iterrows():
            lines.append(
                f"- Sentence {int(row['sentence_id'])}: **{row['trigger']}** "
                f"({row['type']}) -> {row['presupposition']}"
            )

    return "\n".join(lines)


st.sidebar.title("🔑 API Configuration")
entered_key = st.sidebar.text_input(
    "Groq API Key",
    value=os.environ.get("GROQ_API_KEY", ""),
    type="password",
    placeholder="gsk_...",
)
if entered_key:
    os.environ["GROQ_API_KEY"] = entered_key
    st.sidebar.success("✅ LLM verification enabled")
else:
    st.sidebar.info("Running in rule-based mode (no key)")

st.sidebar.markdown("Get a free key at [console.groq.com](https://console.groq.com)")

with st.sidebar.expander("ℹ️ About PragmaLens"):
    st.write(
        "PragmaLens is a computational linguistics tool for detecting presuppositions: "
        "hidden assumptions embedded in ordinary language. It uses spaCy dependency "
        "parsing and formal trigger rules as the core detector, then can use an LLM to "
        "verify and explain the results. The goal is to make media literacy concrete, "
        "inspectable, and fast enough for live analysis."
    )

with st.sidebar.expander("🔬 Trigger Types"):
    for trigger_type, color in TRIGGER_COLORS.items():
        label = TRIGGER_LABELS.get(trigger_type, trigger_type)
        if trigger_type in {"factive", "implicative"}:
            label = f"{label} Verb"
        sample = TRIGGER_SAMPLES.get(trigger_type, "example triggers...")
        st.markdown(
            f"<span style='color:{color};font-size:1.1rem;'>●</span> "
            f"<strong>{html.escape(label)}</strong> — \"{html.escape(sample)}\"",
            unsafe_allow_html=True,
        )


st.markdown('<div class="pl-header">', unsafe_allow_html=True)
st.title("PragmaLens")
st.markdown(
    '<div class="pl-caption">Revealing hidden assumptions in language with formal pragmatics, spaCy, and optional LLM verification.</div>',
    unsafe_allow_html=True,
)
st.markdown("</div>", unsafe_allow_html=True)

overview_tab, impact_tab, analyze_tab, batch_tab, article_tab, eval_tab, prep_tab = st.tabs(
    ["Overview", "🌍 Impact", "Analyze", "Batch Mode", "Article Lab", "Evaluation", "Demo Prep"]
)

with overview_tab:
    st.markdown(
        """
        <section class="hero-banner">
            <h2>Language has a hidden layer.</h2>
            <p>
                Every sentence can carry background assumptions that listeners accept silently.
                PragmaLens makes that layer visible — for the first time, in real time.
            </p>
            <span class="pill-badge">📖 Grounded in formal pragmatics</span>
            <span class="pill-badge">🔬 spaCy + LLM hybrid pipeline</span>
        </section>
        """,
        unsafe_allow_html=True,
    )

    st.subheader("See it in action")
    hero_cols = st.columns([4, 1], gap="large")
    with hero_cols[0]:
        hero_input_text = st.text_input(
            "Try any sentence:",
            key="hero_input",
            value="When did politicians stop caring about ordinary people?",
        )
    with hero_cols[1]:
        st.write("")
        hero_analyze = st.button("→ Analyze")

    if hero_analyze:
        with st.spinner("Checking for hidden assumptions..."):
            hero_results = run_pipeline(hero_input_text, use_llm=has_api_key())
        if hero_results:
            display_presupposition_card(hero_results[0])
        else:
            st.info("No presuppositions detected for this sentence.")

    st.divider()

    st.subheader("Project Summary")
    summary_cols = st.columns([1.2, 1])
    with summary_cols[0]:
        st.markdown(
            """
            <div class="pl-callout">
            <strong>Problem:</strong> people often accept hidden assumptions because wording
            makes them feel like background facts rather than claims to inspect.
            <br><br>
            <strong>Solution:</strong> PragmaLens detects presupposition triggers, extracts
            the implied assumption, highlights the trigger, and explains why the assumption matters.
            <br><br>
            <strong>Method:</strong> the detector uses dependency parsing and formal
            pragmatics rules first. An LLM is then used as a verifier and explanation layer,
            not as the primary detector.
            </div>
            """,
            unsafe_allow_html=True,
        )
    with summary_cols[1]:
        st.markdown(
            """
            <div class="pl-callout">
            <strong>What it uses</strong><br>
            <span class="pl-tag">Streamlit UI</span>
            <span class="pl-tag">spaCy parsing</span>
            <span class="pl-tag">Dependency trees</span>
            <span class="pl-tag">Rule-based NLP</span>
            <span class="pl-tag">Groq/OpenAI/Anthropic optional</span>
            <span class="pl-tag">Plotly visuals</span>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.subheader("Method Visual")
    display_pipeline_visual()

    st.subheader("What the LLM Adds")
    llm_cols = st.columns(4)
    llm_cols[0].info("Verifies whether a rule-based candidate is genuine in context.")
    llm_cols[1].info("Rewrites technical candidates into plain English.")
    llm_cols[2].info("Explains the reader impact and why the wording matters.")
    llm_cols[3].info("Rates subtlety so judges can see which assumptions are hardest to notice.")

    with st.expander("Why this is stronger than a generic chatbot"):
        st.write(
            "The app does not ask an LLM to guess from scratch. It first uses spaCy to find "
            "grammatical evidence: lemmas, part-of-speech tags, dependency labels, subjects, "
            "objects, complements, and clause structure. The LLM only receives structured "
            "candidates from that engine, then verifies and explains them."
        )

with impact_tab:
    st.markdown(
        """
        <section style="background:#EFF6FF;border-left:6px solid #2563EB;border-radius:12px;
                        padding:20px 24px;margin-bottom:24px;">
            <div style="font-size:1.65rem;font-weight:900;color:#0F172A;line-height:1.25;">
                1 in 3 news headlines contains at least one presupposition.
            </div>
            <div style="font-size:1.05rem;color:#475569;margin-top:6px;">
                Most readers never notice. PragmaLens changes that.
            </div>
        </section>
        """,
        unsafe_allow_html=True,
    )

    st.subheader("Who Benefits")
    persona_cols = st.columns(3, gap="large")
    with persona_cols[0]:
        with st.container():
            st.markdown(
                """
                <div class="pl-callout persona-card">
                    <h4>📰 Fact-Checkers & Journalists</h4>
                    <p>
                        When covering political statements or press releases, presuppositions
                        often go unchallenged because they are never stated as claims. PragmaLens
                        surfaces them so journalists can examine what is being assumed, not just asserted.
                    </p>
                    <div class="example-use">Paste a press release → identify all embedded assumptions before writing.</div>
                </div>
                """,
                unsafe_allow_html=True,
            )
    with persona_cols[1]:
        with st.container():
            st.markdown(
                """
                <div class="pl-callout persona-card">
                    <h4>🎓 Teachers & Students</h4>
                    <p>
                        Media literacy curricula teach students to question sources, but rarely
                        teach them to question the grammar itself. PragmaLens teaches the linguistic
                        layer — how presuppositions differ from assertions, and why that matters.
                    </p>
                    <div class="example-use">Use batch mode on a set of exam debate texts → compare assumption loads.</div>
                </div>
                """,
                unsafe_allow_html=True,
            )
    with persona_cols[2]:
        with st.container():
            st.markdown(
                """
                <div class="pl-callout persona-card">
                    <h4>🗳️ Citizens & Policy Researchers</h4>
                    <p>
                        Political advertising and campaign speeches are built on presuppositions.
                        A question like 'When will you fix the crisis you caused?' embeds an accusation
                        before the listener can evaluate whether a crisis exists or who caused it.
                    </p>
                    <div class="example-use">Paste any political statement → see what you are being asked to assume.</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

    st.subheader("Before PragmaLens / After PragmaLens")
    before_cols = st.columns(2, gap="large")
    with before_cols[0]:
        st.markdown("**Without PragmaLens**")
        st.write("When did politicians stop caring about ordinary people?")
        st.markdown(
            """
            <div style="background:#F1F5F9;border:1px solid #CBD5E1;border-radius:8px;
                        padding:14px 16px;color:#334155;line-height:1.55;">
                <strong>Reader sees:</strong> A question about politicians.<br>
                <strong>Reader does not see:</strong> The assumption that politicians did care before,
                and that they have stopped.
            </div>
            """,
            unsafe_allow_html=True,
        )
    with before_cols[1]:
        st.markdown("**With PragmaLens**")
        st.markdown(
            """
            <div style="font-size:1.05rem;line-height:1.8;color:#0F172A;">
                When did politicians
                <mark style="background:#FFEDD5;border-bottom:3px solid #EA580C;
                             color:#0F172A;padding:2px 4px;border-radius:4px;font-weight:800;">
                    stop
                </mark>
                caring about ordinary people?
            </div>
            <div style="margin-top:12px;">
                <span class="compare-presup-chip">✓ Politicians previously cared about ordinary people.</span><br>
                <span class="compare-presup-chip">✓ Politicians have now stopped caring.</span>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.subheader("Where Presuppositions Are Used Against You")
    app_row_one = st.columns(2, gap="large")
    app_row_one[0].info(
        "🏛️ Politics\n\nLoaded questions in debates presuppose guilt, failure, or wrongdoing. "
        "'Will you finally address the corruption in your department?' assumes corruption exists."
    )
    app_row_one[1].info(
        "📺 Advertising\n\nProduct copy uses presuppositions to make claims feel like facts. "
        "'When you're tired of your old phone...' assumes you are tired of it."
    )
    app_row_two = st.columns(2, gap="large")
    app_row_two[0].info(
        "⚖️ Legal Language\n\nLeading questions in courtrooms presuppose the answer. "
        "'Why were you driving so fast?' assumes speeding occurred."
    )
    app_row_two[1].info(
        "🏥 Health Communication\n\nPublic health messaging can use presuppositions to shape behavior. "
        "'Now that we know vaccines cause side effects...' presupposes an established fact."
    )

    st.subheader("Scale of the Problem")
    metric_cols = st.columns(3, gap="large")
    metric_cols[0].metric("Languages with presuppositions", "All of them")
    metric_cols[1].metric("Trigger types detected", "7")
    metric_cols[2].metric("Potential users", "Anyone who reads news")
    st.write(
        "Presuppositions are not a bug in language — they are a feature. Every language "
        "uses them. The problem is not that speakers presuppose things. The problem is "
        "that listeners are rarely taught to notice when it is happening. PragmaLens "
        "is the first practical tool to surface this layer in real time."
    )

with analyze_tab:
    col_input, col_info = st.columns([3, 1], gap="large")

    with col_input:
        example_category = st.selectbox("Try an example", list(EXAMPLES.keys()))
        example = st.selectbox("Example sentence", EXAMPLES[example_category], label_visibility="collapsed")

        user_input = st.text_area(
            "Or enter your own sentence",
            value=example,
            height=115,
            placeholder="Type one sentence to inspect its hidden assumptions.",
        )

        settings_cols = st.columns([1, 2])
        with settings_cols[0]:
            use_llm = st.toggle("Use LLM verification", value=has_api_key())
        with settings_cols[1]:
            if use_llm and not has_api_key():
                st.caption("No API key found. Add GROQ_API_KEY, OPENAI_API_KEY, or ANTHROPIC_API_KEY to .env.")
            elif use_llm:
                st.caption(f"{source_label(use_llm)} enabled.")
            else:
                st.caption("Rule-based analysis only.")

        analyze_btn = st.button("Analyze Presuppositions", type="primary", width="stretch")

    with col_info:
        st.info(
            "A presupposition is a background assumption embedded in language. "
            "It is not the main claim, but listeners often accept it for the sentence to make sense."
        )
        st.markdown(
            """
            <div class="pl-callout">
            <strong>Detected trigger families</strong><br>
            Factive verbs<br>
            Change-of-state verbs<br>
            Implicatives<br>
            Iteratives<br>
            Definite noun phrases<br>
            Clefts<br>
            Temporal clauses
            </div>
            """,
            unsafe_allow_html=True,
        )

    if "last_input" not in st.session_state:
        st.session_state.last_input = ""
    if "last_results" not in st.session_state:
        st.session_state.last_results = []
    if "last_error" not in st.session_state:
        st.session_state.last_error = ""

    if analyze_btn:
        with st.spinner("Analyzing linguistic structure..."):
            try:
                st.session_state.last_results = run_pipeline(user_input, use_llm=use_llm)
                st.session_state.last_input = user_input
                st.session_state.last_error = ""
            except Exception as exc:
                st.session_state.last_results = []
                st.session_state.last_input = user_input
                st.session_state.last_error = str(exc)

    if st.session_state.last_input:
        st.divider()
        st.subheader("Results")

        if st.session_state.last_error:
            st.error(st.session_state.last_error)
        elif not st.session_state.last_results:
            st.success("No presuppositions detected. This sentence is relatively transparent.")
            display_media_literacy_score(compute_media_literacy_score([]))
        else:
            results = st.session_state.last_results
            display_result_summary(results)
            display_media_literacy_score(compute_media_literacy_score(results))
            display_highlighted_text(st.session_state.last_input, results)
            display_legend()

            chart_cols = st.columns(2)
            with chart_cols[0]:
                st.markdown("**Trigger Distribution**")
                display_trigger_chart(results)
            with chart_cols[1]:
                st.markdown("**Confidence by Trigger**")
                display_confidence_chart(results)

            st.markdown("**Presupposition Cards**")
            for presupposition in results:
                display_presupposition_card(presupposition)
                display_rule_evidence(presupposition)
                with st.expander(f"Negation test for '{presupposition.trigger_word}'"):
                    st.write(negation_probe_text(presupposition))

            display_downloads(st.session_state.last_input, results)

    st.divider()
    st.subheader("⚡ Compare Two Texts")
    st.caption(
        "Paste two headlines or statements covering the same topic. See which one embeds more hidden assumptions."
    )

    compare_options = [pair["name"] for pair in COMPARE_PAIRS] + ["Custom"]
    selected_pair_name = st.selectbox("Load example pair", compare_options)
    selected_pair = next((pair for pair in COMPARE_PAIRS if pair["name"] == selected_pair_name), None)
    if selected_pair:
        st.info(
            f"**Newsroom use case:** {selected_pair.get('benefit', 'Compare two versions before publication.')} "
            f"Topic: {selected_pair.get('topic', 'general framing')}."
        )

    if "compare_pair_loaded" not in st.session_state:
        st.session_state.compare_pair_loaded = selected_pair_name
        if selected_pair:
            st.session_state.compare_a = selected_pair["a"]
            st.session_state.compare_b = selected_pair["b"]
        else:
            st.session_state.setdefault("compare_a", "")
            st.session_state.setdefault("compare_b", "")
    elif selected_pair_name != st.session_state.compare_pair_loaded:
        st.session_state.compare_pair_loaded = selected_pair_name
        if selected_pair:
            st.session_state.compare_a = selected_pair["a"]
            st.session_state.compare_b = selected_pair["b"]

    compare_cols = st.columns(2, gap="large")
    with compare_cols[0]:
        text_a = st.text_area("Text A", height=90, key="compare_a")
    with compare_cols[1]:
        text_b = st.text_area("Text B", height=90, key="compare_b")

    if st.button("⚡ Compare Presupposition Load"):
        with st.spinner("Comparing presupposition load..."):
            results_a = run_pipeline(text_a, use_llm=use_llm)
            results_b = run_pipeline(text_b, use_llm=use_llm)

        result_cols = st.columns(2, gap="large")
        with result_cols[0]:
            st.markdown("**Text A**")
            display_highlighted_text(text_a, results_a)
            st.markdown(f"**{len(results_a)} presuppositions detected**")
            for presupposition in results_a:
                display_compact_presupposition(presupposition)
        with result_cols[1]:
            st.markdown("**Text B**")
            display_highlighted_text(text_b, results_b)
            st.markdown(f"**{len(results_b)} presuppositions detected**")
            for presupposition in results_b:
                display_compact_presupposition(presupposition)

        compare_df = pd.DataFrame(
            {
                "Text": ["Text A", "Text B"],
                "Count": [len(results_a), len(results_b)],
            }
        )
        compare_fig = px.bar(
            compare_df,
            x="Text",
            y="Count",
            color="Text",
            color_discrete_map={"Text A": "#4A90E2", "Text B": "#E67E22"},
            text="Count",
            title="Presupposition Count Comparison",
        )
        compare_fig.update_layout(height=300, showlegend=False)
        st.plotly_chart(compare_fig, width="stretch")

        if len(results_a) > len(results_b):
            st.warning(
                f"Text A carries a heavier presupposition load ({len(results_a)} vs {len(results_b)}). "
                "Its framing embeds more assumptions that readers accept without questioning."
            )
        elif len(results_b) > len(results_a):
            st.warning(
                f"Text B carries a heavier presupposition load ({len(results_b)} vs {len(results_a)}). "
                "Its framing embeds more assumptions that readers accept without questioning."
            )
        else:
            st.success(f"Both texts carry a similar presupposition load ({len(results_a)} each).")

with batch_tab:
    st.subheader("Batch Headline and Sentence Analysis")
    st.write(
        "Paste one sentence per line to compare how many hidden assumptions appear across a set of headlines, claims, or debate lines."
    )
    batch_text = st.text_area(
        "Sentences",
        value="\n".join(
            [
                "New report confirms the damage done by previous policies.",
                "When did politicians stop caring about ordinary people?",
                "The agency finally stopped hiding the inspection failures.",
                "She managed to finish before the deadline.",
            ]
        ),
        height=150,
    )
    batch_use_llm = st.toggle("Use LLM in batch mode", value=False)
    if st.button("Run Batch Analysis", type="primary"):
        sentences = [line.strip() for line in batch_text.splitlines() if line.strip()]
        if not sentences:
            st.warning("Add at least one sentence before running batch analysis.")
        else:
            progress = st.progress(0)
            summary_df, detail_df, _ = analyze_sentence_collection(sentences, batch_use_llm, progress)
            st.markdown("**Summary**")
            st.dataframe(summary_df, width="stretch", hide_index=True)
            if not summary_df.empty:
                avg_score = round(summary_df["media_literacy_score"].mean(), 1)
                st.metric("Average Manipulation Load", f"{avg_score}/10")
                fig = px.bar(summary_df, x="sentence", y="count", text="count", title="Presuppositions per Sentence")
                fig.update_layout(height=360, margin=dict(l=10, r=10, t=45, b=10), xaxis_tickangle=-25)
                st.plotly_chart(fig, width="stretch")
            st.markdown("**Detailed Results**")
            st.dataframe(detail_df, width="stretch", hide_index=True)
            st.download_button(
                "Download batch CSV",
                detail_df.to_csv(index=False),
                file_name="pragmalens-batch.csv",
                mime="text/csv",
                width="stretch",
            )

with article_tab:
    st.subheader("Article & Document Audit")
    st.write(
        "Upload a Word document or paste a full article draft. PragmaLens scans sentence by sentence, ranks the most loaded lines, and exports an assumption ledger for editors."
    )

    workflow_cols = st.columns(3, gap="large")
    workflow_cols[0].markdown(
        """
        <div class="workflow-card">
            <strong>1. Bring a draft</strong><br>
            Paste an article, policy brief, transcript, or upload a .docx/.txt/.md file.
        </div>
        """,
        unsafe_allow_html=True,
    )
    workflow_cols[1].markdown(
        """
        <div class="workflow-card">
            <strong>2. Find framing risk</strong><br>
            The app scores each sentence and surfaces hidden assumptions before publication.
        </div>
        """,
        unsafe_allow_html=True,
    )
    workflow_cols[2].markdown(
        """
        <div class="workflow-card">
            <strong>3. Export the ledger</strong><br>
            Download a CSV or editor brief that documents every trigger and assumption.
        </div>
        """,
        unsafe_allow_html=True,
    )

    sample_options = ["Blank draft"] + [item["name"] for item in ARTICLE_EXAMPLES]
    selected_article_sample = st.selectbox("Load newsroom article sample", sample_options)
    selected_article = next(
        (item for item in ARTICLE_EXAMPLES if item["name"] == selected_article_sample),
        None,
    )

    if "article_sample_loaded" not in st.session_state:
        st.session_state.article_sample_loaded = selected_article_sample
        st.session_state.article_text = selected_article["text"] if selected_article else ""
        st.session_state.article_title = selected_article_sample if selected_article else "Untitled draft"
    elif selected_article_sample != st.session_state.article_sample_loaded:
        st.session_state.article_sample_loaded = selected_article_sample
        st.session_state.article_text = selected_article["text"] if selected_article else ""
        st.session_state.article_title = selected_article_sample if selected_article else "Untitled draft"

    uploaded_article = st.file_uploader(
        "Upload an article or document",
        type=["docx", "txt", "md"],
        help="Supports Word documents, plain text, and Markdown drafts.",
    )
    if uploaded_article is not None and uploaded_article.name != st.session_state.get("article_upload_name"):
        try:
            st.session_state.article_text = extract_text_from_upload(uploaded_article)
            st.session_state.article_title = uploaded_article.name
            st.session_state.article_upload_name = uploaded_article.name
            st.success(f"Loaded {uploaded_article.name}")
        except RuntimeError as exc:
            st.error(str(exc))

    title_cols = st.columns([2, 1, 1], gap="large")
    with title_cols[0]:
        article_title = st.text_input("Audit title", key="article_title")
    with title_cols[1]:
        sentence_limit = st.slider("Sentence limit", min_value=5, max_value=80, value=35, step=5)
    with title_cols[2]:
        article_use_llm = st.toggle("Use LLM", value=False, key="article_use_llm")

    article_text = st.text_area(
        "Article, draft, transcript, or policy text",
        key="article_text",
        height=280,
        placeholder="Paste a full article draft or upload a document above.",
    )

    if article_use_llm and not has_api_key():
        st.caption("No API key found. Article Lab will fall back to rule-based output.")
    elif article_use_llm:
        st.caption(f"{source_label(article_use_llm)} enabled for article audit.")
    else:
        st.caption("Rule-based article audit. Recommended for fast whole-document scans.")

    if st.button("Audit Article / Document", type="primary", width="stretch"):
        sentences = split_text_into_sentences(article_text, sentence_limit)
        if not sentences:
            st.warning("No analyzable sentences found. Add more article text or upload a supported document.")
        else:
            st.info(f"Scanning {len(sentences)} sentences.")
            progress = st.progress(0)
            summary_df, detail_df, sentence_results = analyze_sentence_collection(
                sentences,
                article_use_llm,
                progress,
            )
            st.session_state.article_audit = {
                "title": article_title,
                "summary_df": summary_df,
                "detail_df": detail_df,
                "sentence_results": sentence_results,
            }

    article_audit = st.session_state.get("article_audit")
    if article_audit:
        summary_df = article_audit["summary_df"]
        detail_df = article_audit["detail_df"]
        sentence_results = article_audit["sentence_results"]
        all_presuppositions = [
            item
            for sentence_result in sentence_results
            for item in sentence_result["results"]
        ]

        st.divider()
        st.subheader("Article Audit Results")
        if summary_df.empty:
            st.info("No analyzable sentences were found.")
        else:
            total_presuppositions = int(summary_df["count"].sum())
            avg_load = round(summary_df["media_literacy_score"].mean(), 1)
            high_load_count = int((summary_df["media_literacy_score"] > 6).sum())
            trigger_family_count = len({item.trigger_type for item in all_presuppositions})

            metric_cols = st.columns(4, gap="large")
            metric_cols[0].metric("Sentences scanned", len(summary_df))
            metric_cols[1].metric("Total presuppositions", total_presuppositions)
            metric_cols[2].metric("Average load", f"{avg_load}/10")
            metric_cols[3].metric("Trigger families", trigger_family_count)

            if high_load_count:
                st.warning(
                    f"{high_load_count} sentence(s) scored above 6/10. Review these lines for loaded framing before publication."
                )
            else:
                st.success("No high-load sentences found. The draft is relatively transparent by the current rule set.")

            chart_df = summary_df.assign(sentence_label=summary_df["sentence_id"].map(lambda value: f"S{value}"))
            load_fig = px.bar(
                chart_df,
                x="sentence_label",
                y="media_literacy_score",
                color="score_label",
                color_discrete_map={
                    "Transparent": "#22C55E",
                    "Low Load": "#22C55E",
                    "Moderate Load": "#F59E0B",
                    "High Load": "#EF4444",
                },
                hover_data=["sentence", "count", "trigger_types"],
                text="media_literacy_score",
                title="Framing Load Across the Document",
            )
            load_fig.update_layout(
                height=340,
                margin=dict(l=10, r=10, t=45, b=10),
                yaxis=dict(range=[0, 10], title="Load score"),
                xaxis=dict(title="Sentence"),
            )
            st.plotly_chart(load_fig, width="stretch")

            if all_presuppositions:
                st.markdown("**Trigger Distribution Across Document**")
                display_trigger_chart(all_presuppositions)

            top_df = summary_df.sort_values(
                ["media_literacy_score", "count", "avg_subtlety"],
                ascending=False,
            ).head(10)
            st.markdown("**Top Framing Risks**")
            st.dataframe(
                top_df[
                    [
                        "sentence_id",
                        "media_literacy_score",
                        "count",
                        "trigger_types",
                        "sentence",
                    ]
                ],
                width="stretch",
                hide_index=True,
            )

            st.markdown("**High-Risk Sentence Review**")
            sentence_lookup = {item["sentence_id"]: item for item in sentence_results}
            for _, row in top_df.head(5).iterrows():
                sentence_result = sentence_lookup[int(row["sentence_id"])]
                with st.expander(
                    f"Sentence {int(row['sentence_id'])} · {row['media_literacy_score']}/10 · {int(row['count'])} presuppositions"
                ):
                    st.markdown(
                        f"""
                        <div class="risk-sentence">
                            <strong>Original sentence</strong><br>
                            {html.escape(sentence_result["sentence"])}
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                    display_highlighted_text(sentence_result["sentence"], sentence_result["results"])
                    if sentence_result["results"]:
                        for presupposition in sentence_result["results"]:
                            display_compact_presupposition(presupposition)
                    else:
                        st.info("No presuppositions detected in this sentence.")

            download_cols = st.columns(2, gap="large")
            download_cols[0].download_button(
                "Download assumption ledger CSV",
                detail_df.to_csv(index=False),
                file_name="pragmalens-assumption-ledger.csv",
                mime="text/csv",
                width="stretch",
                disabled=detail_df.empty,
            )
            download_cols[1].download_button(
                "Download editor brief",
                build_document_audit_report(article_audit["title"], summary_df, detail_df),
                file_name="pragmalens-editor-brief.md",
                mime="text/markdown",
                width="stretch",
            )

with eval_tab:
    st.subheader("Internal Benchmark")
    st.write(
        "This lightweight benchmark checks whether the rule engine detects the expected trigger families for curated examples."
    )
    if st.button("Run Evaluation", type="primary"):
        eval_df = run_evaluation()
        pass_rate = eval_df["pass"].mean() if not eval_df.empty else 0
        metric_cols = st.columns(3)
        metric_cols[0].metric("Cases", len(eval_df))
        metric_cols[1].metric("Pass rate", f"{pass_rate:.0%}")
        metric_cols[2].metric("Total detections", int(eval_df["count"].sum()))
        st.dataframe(eval_df, width="stretch", hide_index=True)
        chart_df = eval_df.assign(result=eval_df["pass"].map({True: "Pass", False: "Fail"}))
        fig = px.histogram(chart_df, x="result", color="result", title="Evaluation Results")
        fig.update_layout(height=280, margin=dict(l=10, r=10, t=45, b=10), showlegend=False)
        st.plotly_chart(fig, width="stretch")

with prep_tab:
    st.subheader("Demo Sequence")
    for item in KILLER_DEMOS:
        with st.expander(item["title"]):
            st.code(item["sentence"], language=None)
            st.write(item["why"])

    st.subheader("60-Second Pitch")
    st.markdown(
        """
        PragmaLens reveals hidden assumptions in language. It solves the problem that
        loaded questions, headlines, and everyday phrasing can make readers accept a
        background claim without noticing. The system uses spaCy dependency parsing and
        formal presupposition trigger rules to detect candidates, then optionally uses
        an LLM to verify and explain the result in plain English.
        """
    )

    st.subheader("Live Demo Flow")
    st.markdown(
        """
        1. Start with the quick demo in Overview to show the hidden assumption layer.
        2. Open Analyze and run the Media Literacy Score demo sentence.
        3. Use Compare Two Texts to show how two headlines about the same event carry different framing loads.
        4. Finish in Article Lab by loading a newsroom draft and exporting the assumption ledger.
        """
    )

st.caption("Built at LingHacks VII | Grounded in formal pragmatics theory")
